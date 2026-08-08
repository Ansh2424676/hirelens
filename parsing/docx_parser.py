from docx import Document


def extract_text_from_docx(file_path: str) -> dict:
    """Extract paragraphs and table text from a DOCX resume."""
    try:
        document = Document(file_path)
        sections = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                sections.append(text)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                row_text = " | ".join(cell for cell in cells if cell)
                if row_text:
                    sections.append(row_text)

        extracted_text = "\n".join(sections).strip()

        if not extracted_text:
            return {
                "success": False,
                "text": "",
                "error": "No readable text was found in this DOCX file.",
            }

        return {
            "success": True,
            "text": extracted_text,
            "error": None,
        }

    except Exception as exc:
        return {
            "success": False,
            "text": "",
            "error": f"Unable to read the DOCX file: {exc}",
        }